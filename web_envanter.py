import streamlit as st
import sqlite3
import pandas as pd
import datetime
import matplotlib.pyplot as plt
import io
import docx

# --- SAYFA AYARLARI ---
st.set_page_config(page_title="ABB Afet Lojistik Portalı", layout="wide")
DB_YOLU = 'afet_lojistik_web.db'

# --- YARDIMCI FONKSİYONLAR (EXCEL & WORD ÜRETİMİ) ---
@st.cache_data
def df_to_excel(df, sheet_name="Sayfa1"):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return output.getvalue()

@st.cache_data
def df_to_word(df, title):
    doc = docx.Document()
    doc.add_heading(title, 0)
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    t.style = 'Table Grid'
    for j in range(df.shape[-1]):
        t.cell(0,j).text = str(df.columns[j])
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def detayli_rapor_excel(df_stok, df_malzeme, df_yer, detayli_mi):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_stok.to_excel(writer, index=False, sheet_name='Güncel Stok')
        if detayli_mi:
            df_malzeme.to_excel(writer, index=False, sheet_name='Malzeme Analizi')
            df_yer.to_excel(writer, index=False, sheet_name='Bölge Analizi')
    return output.getvalue()

def detayli_rapor_word(df_stok, df_malzeme, df_yer, detayli_mi):
    doc = docx.Document()
    doc.add_heading('Afet Eğitim ve Lojistik Merkezi - Güncel Envanter Raporu', 0)
    
    t_stok = doc.add_table(df_stok.shape[0]+1, df_stok.shape[1])
    t_stok.style = 'Table Grid'
    for j in range(df_stok.shape[-1]): t_stok.cell(0,j).text = str(df_stok.columns[j])
    for i in range(df_stok.shape[0]):
        for j in range(df_stok.shape[-1]): t_stok.cell(i+1,j).text = str(df_stok.values[i,j])
            
    if detayli_mi:
        doc.add_heading('Saha Operasyon Analizi - En Çok Giden Malzemeler', 1)
        t_malz = doc.add_table(df_malzeme.shape[0]+1, df_malzeme.shape[1])
        t_malz.style = 'Table Grid'
        for j in range(df_malzeme.shape[-1]): t_malz.cell(0,j).text = str(df_malzeme.columns[j])
        for i in range(df_malzeme.shape[0]):
            for j in range(df_malzeme.shape[-1]): t_malz.cell(i+1,j).text = str(df_malzeme.values[i,j])
                
        doc.add_heading('Saha Operasyon Analizi - Bölgeler', 1)
        t_yer = doc.add_table(df_yer.shape[0]+1, df_yer.shape[1])
        t_yer.style = 'Table Grid'
        for j in range(df_yer.shape[-1]): t_yer.cell(0,j).text = str(df_yer.columns[j])
        for i in range(df_yer.shape[0]):
            for j in range(df_yer.shape[-1]): t_yer.cell(i+1,j).text = str(df_yer.values[i,j])

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# --- VERİTABANI BAŞLATMA ---
def init_db():
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS malzemeler (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        malzeme_adi TEXT UNIQUE NOT NULL,
                        kategori TEXT NOT NULL,
                        birim TEXT NOT NULL,
                        toplam_stok REAL DEFAULT 0,
                        sahadaki_stok REAL DEFAULT 0,
                        zayi_stok REAL DEFAULT 0)''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS hareketler (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        malzeme_adi TEXT NOT NULL,
                        islem_tipi TEXT NOT NULL,
                        miktar REAL NOT NULL,
                        tarih TEXT NOT NULL,
                        yer TEXT NOT NULL)''')
    conn.commit()
    conn.close()

init_db()

def get_data(query, params=()):
    conn = sqlite3.connect(DB_YOLU)
    df = pd.read_sql_query(query, conn, params=params)
    conn.close()
    return df

def run_query(query, params=()):
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute(query, params)
    conn.commit()
    conn.close()

# --- KURUMSAL BAŞLIK VE LOGOLAR ---
col_logo1, col_baslik, col_logo2 = st.columns([1, 4, 1])
with col_logo1:
    try: st.image("abb.JPG", use_column_width=True)
    except: st.warning("ABB Logosu eksik.")
with col_baslik:
    st.markdown("<h3 style='text-align: center; color: #555;'>Afet İşleri ve Risk Yönetimi Dairesi Başkanlığı</h3>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center; color: #2C3E50; font-weight: bold;'>Afet Eğitim ve Lojistik Merkezi</h2>", unsafe_allow_html=True)
with col_logo2:
    try: st.image("db.JPG", use_column_width=True)
    except: st.warning("Daire Logosu eksik.")
st.markdown("---")

tab1, tab2, tab3 = st.tabs(["📦 1. Malzeme Tanımlama & Düzenleme", "🚚 2. Depo İşlemleri (Sevk/Zayi)", "📊 3. Analiz ve Raporlar"])

# ================= SEKME 1: MALZEME TANIMLAMA =================
with tab1:
    col1, col2 = st.columns([1, 2])
    with col1:
        st.write("### Yeni Malzeme Ekle")
        with st.form("malzeme_ekle_form", clear_on_submit=True):
            kategori = st.selectbox("Kategori", ["Çadır & Barınma", "Konteyner", "Isıtıcı & Jeneratör", "Mobil Araçlar", "Hijyen & Tuvalet", "Arama Kurtarma Ekipmanı", "Diğer"])
            malzeme_adi = st.text_input("Malzeme Adı")
            birim = st.selectbox("Birim", ["Adet", "Litre", "Kg", "Ton", "Kutu", "Koli"])
            submit = st.form_submit_button("Sisteme Kaydet")
            if submit and malzeme_adi:
                try:
                    run_query("INSERT INTO malzemeler (malzeme_adi, kategori, birim, toplam_stok, sahadaki_stok, zayi_stok) VALUES (?, ?, ?, 0, 0, 0)", (malzeme_adi, kategori, birim))
                    st.success(f"{malzeme_adi} eklendi!")
                    st.rerun()
                except: st.error("Bu malzeme zaten var!")
                    
        st.write("### Malzeme Güncelle / Sil")
        malzemeler_df = get_data("SELECT id, malzeme_adi FROM malzemeler ORDER BY malzeme_adi")
        if not malzemeler_df.empty:
            islem_turu = st.radio("İşlem:", ["Adını Güncelle", "Tamamen Sil"])
            secilen_malzeme = st.selectbox("Malzeme Seç:", malzemeler_df['malzeme_adi'].tolist())
            if islem_turu == "Tamamen Sil":
                if st.button("Kalıcı Olarak Sil", type="primary"):
                    run_query("DELETE FROM malzemeler WHERE malzeme_adi = ?", (secilen_malzeme,))
                    run_query("DELETE FROM hareketler WHERE malzeme_adi = ?", (secilen_malzeme,))
                    st.success("Silindi.")
                    st.rerun()
            else:
                yeni_ad = st.text_input("Yeni Ad:", value=secilen_malzeme)
                if st.button("Güncelle", type="primary"):
                    if yeni_ad and yeni_ad != secilen_malzeme:
                        run_query("UPDATE malzemeler SET malzeme_adi = ? WHERE malzeme_adi = ?", (yeni_ad, secilen_malzeme))
                        run_query("UPDATE hareketler SET malzeme_adi = ? WHERE malzeme_adi = ?", (yeni_ad, secilen_malzeme))
                        st.success("Güncellendi.")
                        st.rerun()

    with col2:
        st.write("### Sistemde Tanımlı Malzemeler")
        st.dataframe(get_data("SELECT id as ID, malzeme_adi as 'Malzeme Adı', kategori as Kategori, birim as Birim FROM malzemeler ORDER BY malzeme_adi"), use_container_width=True, hide_index=True)

# ================= SEKME 2: GİRİŞ / ÇIKIŞ / ZAYİ =================
with tab2:
    col_islem, col_gecmis = st.columns([1, 2])
    with col_islem:
        st.write("### Yeni Depo İşlemi")
        with st.form("islem_form", clear_on_submit=True):
            islemler = ["Sahaya Gönder (Depodan Çıkış)", "Sahadan Döndü (Depoya Sağlam Giriş)", "Yeni Alım (Depoya Ekle)", "Sahada Hurda/Zayi Oldu", "Depoda Hurda/Zayi Oldu"]
            islem_tipi = st.selectbox("İşlem Tipi", islemler)
            malz_list = get_data("SELECT malzeme_adi FROM malzemeler ORDER BY malzeme_adi")['malzeme_adi'].tolist()
            malzeme_secim = st.selectbox("Malzeme Seç", malz_list if malz_list else ["Tanımlı malzeme yok"])
            
            # TAM SAYI OLARAK GÜNCELLENDİ
            miktar = st.number_input("Miktar (Tam Sayı)", min_value=1, step=1, format="%d")
            # TARİH FORMATI GÜNCELLENDİ
            tarih = st.date_input("Tarih", datetime.date.today(), format="DD/MM/YYYY")
            yer = st.text_input("İlgili Yer / Etkinlik")
            islem_kaydet = st.form_submit_button("Kaydet ve Stok Güncelle")
            
            if islem_kaydet and malzeme_secim != "Tanımlı malzeme yok" and yer:
                durum_df = get_data("SELECT toplam_stok, sahadaki_stok, zayi_stok FROM malzemeler WHERE malzeme_adi=?", (malzeme_secim,))
                toplam, sahada, zayi = durum_df.iloc[0]['toplam_stok'], durum_df.iloc[0]['sahadaki_stok'], durum_df.iloc[0]['zayi_stok']
                depoda = toplam - sahada - zayi
                hata = False
                
                if islem_tipi == "Sahaya Gönder (Depodan Çıkış)" and miktar > depoda: st.error(f"Stok yetersiz! Depoda: {int(depoda)}"); hata = True
                elif islem_tipi == "Sahadan Döndü (Depoya Sağlam Giriş)" and miktar > sahada: st.warning(f"Saha stoğu yetersiz: {int(sahada)}"); hata = True
                elif islem_tipi == "Sahada Hurda/Zayi Oldu" and miktar > sahada: st.error(f"Saha stoğu yetersiz: {int(sahada)}"); hata = True
                elif islem_tipi == "Depoda Hurda/Zayi Oldu" and miktar > depoda: st.error(f"Depo stoğu yetersiz: {int(depoda)}"); hata = True
                
                if not hata:
                    if islem_tipi == "Sahaya Gönder (Depodan Çıkış)": run_query("UPDATE malzemeler SET sahadaki_stok = sahadaki_stok + ? WHERE malzeme_adi = ?", (miktar, malzeme_secim))
                    elif islem_tipi == "Sahadan Döndü (Depoya Sağlam Giriş)": run_query("UPDATE malzemeler SET sahadaki_stok = sahadaki_stok - ? WHERE malzeme_adi = ?", (miktar, malzeme_secim))
                    elif islem_tipi == "Yeni Alım (Depoya Ekle)": run_query("UPDATE malzemeler SET toplam_stok = toplam_stok + ? WHERE malzeme_adi = ?", (miktar, malzeme_secim))
                    elif islem_tipi == "Sahada Hurda/Zayi Oldu": run_query("UPDATE malzemeler SET sahadaki_stok = sahadaki_stok - ?, zayi_stok = zayi_stok + ? WHERE malzeme_adi = ?", (miktar, miktar, malzeme_secim))
                    elif islem_tipi == "Depoda Hurda/Zayi Oldu": run_query("UPDATE malzemeler SET zayi_stok = zayi_stok + ? WHERE malzeme_adi = ?", (miktar, malzeme_secim))
                    
                    run_query("INSERT INTO hareketler (malzeme_adi, islem_tipi, miktar, tarih, yer) VALUES (?, ?, ?, ?, ?)", (malzeme_secim, islem_tipi, miktar, tarih.strftime("%d.%m.%Y"), yer))
                    st.success("Kaydedildi!")
                    st.rerun()

    with col_gecmis:
        st.write("### İşlem Defteri")
        df_hareket = get_data("SELECT id as ID, tarih as Tarih, malzeme_adi as Malzeme, islem_tipi as 'İşlem Tipi', miktar as Miktar, yer as 'Görev Yeri' FROM hareketler ORDER BY id DESC")
        
        # Miktarları tam sayı göster
        if not df_hareket.empty: df_hareket['Miktar'] = df_hareket['Miktar'].astype(int)
        
        st.dataframe(df_hareket, use_container_width=True, hide_index=True)
        
        if not df_hareket.empty:
            with st.expander("⚠️ İşlemi Geri Al (Sil)"):
                sil_id = st.selectbox("ID Seçin:", df_hareket['ID'].tolist())
                secili = df_hareket[df_hareket['ID'] == sil_id].iloc[0]
                if st.button("Geri Al", type="primary"):
                    malz, islem, miktar = secili['Malzeme'], secili['İşlem Tipi'], float(secili['Miktar'])
                    if islem == "Sahaya Gönder (Depodan Çıkış)": run_query("UPDATE malzemeler SET sahadaki_stok = sahadaki_stok - ? WHERE malzeme_adi = ?", (miktar, malz))
                    elif islem == "Sahadan Döndü (Depoya Sağlam Giriş)": run_query("UPDATE malzemeler SET sahadaki_stok = sahadaki_stok + ? WHERE malzeme_adi = ?", (miktar, malz))
                    elif islem == "Yeni Alım (Depoya Ekle)": run_query("UPDATE malzemeler SET toplam_stok = toplam_stok - ? WHERE malzeme_adi = ?", (miktar, malz))
                    elif islem == "Sahada Hurda/Zayi Oldu": run_query("UPDATE malzemeler SET sahadaki_stok = sahadaki_stok + ?, zayi_stok = zayi_stok - ? WHERE malzeme_adi = ?", (miktar, miktar, malz))
                    elif islem == "Depoda Hurda/Zayi Oldu": run_query("UPDATE malzemeler SET zayi_stok = zayi_stok - ? WHERE malzeme_adi = ?", (miktar, malz))
                    run_query("DELETE FROM hareketler WHERE id = ?", (int(sil_id),))
                    st.success("Geri alındı!")
                    st.rerun()
            
            # EXCEL VE WORD İNDİRME BUTONLARI
            col_ex, col_wd = st.columns(2)
            with col_ex:
                st.download_button("📊 Excel Olarak İndir", data=df_to_excel(df_hareket, "İşlem Defteri"), file_name='islem_defteri.xlsx', mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            with col_wd:
                st.download_button("📝 Word Olarak İndir", data=df_to_word(df_hareket, "Afet Lojistik - İşlem Defteri"), file_name='islem_defteri.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# ================= SEKME 3: RAPOR VE ANALİZ =================
with tab3:
    st.write("### 🏢 Depolardaki Anlık Genel Durum")
    df_stok = get_data("SELECT malzeme_adi as Malzeme, kategori as Kategori, toplam_stok, sahadaki_stok, zayi_stok, birim as Birim FROM malzemeler")
    df_stok['Depoda (Mevcut)'] = df_stok['toplam_stok'] - df_stok['sahadaki_stok'] - df_stok['zayi_stok']
    df_stok = df_stok.rename(columns={'toplam_stok': 'Toplam Alınan', 'sahadaki_stok': 'Görevde (Sahada)', 'zayi_stok': 'Zayi / Hurda'})
    df_stok = df_stok[['Malzeme', 'Kategori', 'Toplam Alınan', 'Görevde (Sahada)', 'Zayi / Hurda', 'Depoda (Mevcut)', 'Birim']]
    
    # Tam sayı formatlaması
    for col in ['Toplam Alınan', 'Görevde (Sahada)', 'Zayi / Hurda', 'Depoda (Mevcut)']:
        df_stok[col] = df_stok[col].astype(int)
        
    df_stok = df_stok.sort_values(by='Depoda (Mevcut)', ascending=False)
    st.dataframe(df_stok, use_container_width=True, hide_index=True)
    
    # Analiz DataFrameleri
    df_analiz = get_data("SELECT malzeme_adi as Malzeme, SUM(miktar) as 'Toplam Giden', COUNT(id) as 'Sevkiyat Sayısı' FROM hareketler WHERE islem_tipi='Sahaya Gönder (Depodan Çıkış)' GROUP BY malzeme_adi ORDER BY SUM(miktar) DESC")
    df_yer = get_data("SELECT yer as 'Görev Yeri', COUNT(id) as 'Operasyon Sayısı' FROM hareketler WHERE islem_tipi='Sahaya Gönder (Depodan Çıkış)' GROUP BY yer ORDER BY COUNT(id) DESC")
    
    if not df_analiz.empty: df_analiz['Toplam Giden'] = df_analiz['Toplam Giden'].astype(int)

    st.markdown("---")
    
    # DETAYLI RAPOR İNDİRME SEÇENEĞİ VE BUTONLAR
    st.write("#### 📥 Gelişmiş Rapor Çıktısı Al")
    detay_secim = st.checkbox("İndirilecek rapora Saha Operasyon Analizlerini (En çok gidenler vb.) dahil et")
    
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        st.download_button("📊 Raporu Excel İndir", data=detayli_rapor_excel(df_stok, df_analiz, df_yer, detay_secim), file_name='guncel_envanter.xlsx', mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    with col_r2:
        st.download_button("📝 Raporu Word İndir", data=detayli_rapor_word(df_stok, df_analiz, df_yer, detay_secim), file_name='guncel_envanter.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    st.markdown("---")
    col_grafik, col_analiz_tablo = st.columns(2)
    
    with col_grafik:
        st.write("### 📈 Kullanıma Hazır Malzeme Grafiği")
        df_grafik = df_stok[df_stok['Depoda (Mevcut)'] > 0]
        if not df_grafik.empty:
            fig, ax = plt.subplots(figsize=(8, 5))
            bars = ax.bar(df_grafik['Malzeme'], df_grafik['Depoda (Mevcut)'], color='#3498DB')
            
            # GRAFİK BARLARININ ÜZERİNE SAYILARI YAZDIRMA
            ax.bar_label(bars, padding=3, fmt='%d', color='black', fontweight='bold')
            
            plt.xticks(rotation=45, ha='right')
            plt.ylabel('Depo Mevcudu')
            st.pyplot(fig)
        else:
            st.info("Grafik çizilecek mevcut stok bulunamadı.")
            
    with col_analiz_tablo:
        st.write("### 🎯 Saha Operasyon Analizi")
        st.write("**En Çok Sahaya Gönderilen Malzemeler**")
        st.dataframe(df_analiz, use_container_width=True, hide_index=True)
        st.write("**En Çok Sevkiyat Yapılan Bölgeler**")
        st.dataframe(df_yer, use_container_width=True, hide_index=True)

# --- FOOTER ---
st.markdown("---")
st.markdown("<p style='text-align: right; color: gray; font-style: italic;'>Tüm hakları saklıdır © Erdem KUŞÇU - 2026</p>", unsafe_allow_html=True)